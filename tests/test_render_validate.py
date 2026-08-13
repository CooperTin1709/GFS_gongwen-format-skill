from __future__ import annotations

import copy
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from scripts.classify import classify_document
from scripts.docx_utils import (
    FONT_ATTRIBUTES,
    font_values,
    indent_values,
    is_digit_run_text,
    set_first_line_indent_chars,
    split_text_by_digit_runs,
)
from scripts.input_adapter import adapt_source_text
from scripts.render_docx import render_document
from scripts.utils import load_rules, resolved_format_rule
from scripts.validate import validate_document


ROOT = Path(__file__).resolve().parents[1]
TEMP_ROOT = ROOT / ".tmp" / "tests"
SOURCE_TEXT = """关于2026年度工作安排的通知

行领导：

一、2026年度总体要求
项目编号为AI-001，完成率为100%。

（一）第2阶段任务
计划于2026年8月11日前完成。

1. 完成第3轮测试
测试编号001。

（1）处理第4类问题
相关指标为3.14。

附件：1.2026年度任务表
　　　2.测试说明

XX部门
2026年8月11日"""


def build_output(
    folder: Path,
) -> tuple[dict[str, object], Path, dict[str, object], list[dict[str, object]]]:
    canonical = adapt_source_text(SOURCE_TEXT)
    classified, analysis = classify_document(canonical)
    if analysis["review_count"]:
        raise AssertionError(analysis["review_items"])
    output = folder / "formatted.docx"
    _, plan = render_document(classified, output)
    result = validate_document(classified, output)
    return classified, output, result, plan


def paragraph_with_text(document: Document, text: str):
    return next(paragraph for paragraph in document.paragraphs if paragraph.text == text)


def paragraph_with_type(
    document: Document, classified: dict[str, object], paragraph_type: str, occurrence: int = 0
):
    texts = [
        item["text"]
        for item in classified["paragraphs"]
        if item["classification"] == paragraph_type
    ]
    return paragraph_with_text(document, texts[occurrence])


def set_all_run_fonts(run, font_name: str) -> None:
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in FONT_ATTRIBUTES:
        fonts.set(qn(f"w:{attribute}"), font_name)


class RenderValidateTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        TEMP_ROOT.mkdir(parents=True, exist_ok=True)

    def assert_validation_failed(
        self,
        classified: dict[str, object],
        output: Path,
        expected_check: str,
    ) -> None:
        result = validate_document(classified, output)
        self.assertEqual(result["status"], "VALIDATION_FAILED")
        self.assertFalse(result["validation_passed"])
        self.assertIn(expected_check, {error["check"] for error in result["errors"]})

    def test_render_reopen_and_strict_validation(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, result, plan = build_output(Path(directory))
            self.assertEqual(result["status"], "SUCCESS", result["errors"])
            self.assertEqual(result["verified_against"], "browser_extracted_text")
            self.assertTrue(result["validation"]["text_identical"])
            self.assertEqual(result["validation"]["blank_after_title_count"], 1)
            self.assertEqual(result["validation"]["blank_before_attachment_count"], 1)

            reopened = Document(output)
            expected = [item["text"] for item in classified["paragraphs"]]
            actual = [
                paragraph.text
                for paragraph in reopened.paragraphs
                if paragraph.text != ""
            ]
            self.assertEqual(actual, expected)
            self.assertEqual([item["text"] for item in plan], [
                paragraph.text for paragraph in reopened.paragraphs
            ])

            rules = load_rules()
            digit_font = rules["global"]["digit_font"]
            nonempty = [
                paragraph for paragraph in reopened.paragraphs if paragraph.text != ""
            ]
            for item, paragraph in zip(classified["paragraphs"], nonempty):
                paragraph_type = item["classification"]
                rule = resolved_format_rule(rules, paragraph_type)
                self.assertEqual(paragraph.text, item["text"])
                self.assertEqual(
                    [(run.text, is_digit_run_text(run.text)) for run in paragraph.runs],
                    split_text_by_digit_runs(item["text"]),
                )
                for run in paragraph.runs:
                    expected_font = digit_font if is_digit_run_text(run.text) else rule["font"]
                    self.assertEqual(
                        font_values(run),
                        {attribute: expected_font for attribute in FONT_ATTRIBUTES},
                    )
                    self.assertEqual(run.font.size.pt, float(rule["size_pt"]))

                values = indent_values(paragraph)
                expected_chars = float(rule["first_line_indent_chars"])
                self.assertEqual(
                    values["firstLineChars"],
                    str(round(expected_chars * 100)) if expected_chars else None,
                )
                self.assertIsNone(values["firstLine"])
                self.assertIsNone(values["hanging"])
                self.assertIsNone(values["hangingChars"])

            by_text = {paragraph.text: paragraph for paragraph in nonempty}
            self.assertEqual(
                by_text["关于2026年度工作安排的通知"].alignment,
                WD_ALIGN_PARAGRAPH.CENTER,
            )
            self.assertEqual(by_text["行领导："].alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertEqual(by_text["XX部门"].alignment, WD_ALIGN_PARAGRAPH.RIGHT)
            self.assertEqual(
                by_text["2026年8月11日"].alignment, WD_ALIGN_PARAGRAPH.RIGHT
            )

            formats = result["validation"]["observed_formats"]
            self.assertEqual(formats["title"]["size_pt"], [22.0])
            self.assertEqual(formats["title"]["alignment"], ["CENTER"])
            self.assertEqual(formats["salutation"]["alignment"], ["LEFT"])
            self.assertEqual(formats["signature"]["alignment"], ["RIGHT"])
            for paragraph_type in (
                "heading_1",
                "heading_2",
                "heading_3",
                "heading_4",
                "body",
                "attachment",
            ):
                self.assertEqual(
                    formats[paragraph_type]["first_line_indent_chars"], [2.0]
                )
                self.assertEqual(
                    formats[paragraph_type]["line_spacing_pt"], [30.0]
                )
            for paragraph_type in ("title", "salutation", "signature", "blank"):
                self.assertEqual(
                    formats[paragraph_type]["first_line_indent_chars"], [0.0]
                )

    def test_split_text_by_digit_runs_preserves_all_characters(self) -> None:
        text = "本项目２０２６年度编号为AI-001，指标3.14%。"
        parts = split_text_by_digit_runs(text)
        self.assertEqual("".join(part for part, _ in parts), text)
        self.assertEqual(
            [part for part, is_digit in parts if is_digit],
            ["２０２６", "001", "3", "14"],
        )

    def test_validator_rejects_text_change(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            paragraph = paragraph_with_text(
                document, "项目编号为AI-001，完成率为100%。"
            )
            paragraph.runs[2].text = "002"
            document.save(output)
            self.assert_validation_failed(classified, output, "text_integrity")

    def test_validator_rejects_wrong_title_font(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            set_all_run_fonts(document.paragraphs[0].runs[0], "宋体")
            document.save(output)
            self.assert_validation_failed(classified, output, "font_eastAsia")

    def test_validator_rejects_multiple_line_spacing(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            paragraph_with_text(document, "项目编号为AI-001，完成率为100%。").paragraph_format.line_spacing = 1.5
            document.save(output)
            self.assert_validation_failed(classified, output, "line_spacing")

    def test_validator_rejects_deleted_required_blank(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            blank = document.paragraphs[1]._element
            blank.getparent().remove(blank)
            document.save(output)
            self.assert_validation_failed(classified, output, "blank_after_title")

    def test_validator_rejects_extra_blank(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            first_blank = document.paragraphs[1]._element
            first_blank.addnext(copy.deepcopy(first_blank))
            document.save(output)
            self.assert_validation_failed(classified, output, "blank_after_title")

    def test_validator_rejects_wrong_heading_2_size(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            paragraph_with_text(document, "（一）第2阶段任务").runs[0].font.size = Pt(15)
            document.save(output)
            self.assert_validation_failed(classified, output, "font_size")

    def test_validator_rejects_body_without_indent(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            set_first_line_indent_chars(
                paragraph_with_text(document, "项目编号为AI-001，完成率为100%。"),
                0,
            )
            document.save(output)
            self.assert_validation_failed(classified, output, "first_line_indent")

    def test_validator_rejects_indented_salutation(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            set_first_line_indent_chars(paragraph_with_text(document, "行领导："), 2)
            document.save(output)
            self.assert_validation_failed(classified, output, "first_line_indent")

    def test_validator_rejects_left_aligned_signature(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            paragraph_with_text(document, "XX部门").alignment = WD_ALIGN_PARAGRAPH.LEFT
            document.save(output)
            self.assert_validation_failed(classified, output, "alignment")

    def test_validator_rejects_centered_signature(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            paragraph_with_type(
                document, classified, "signature"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.save(output)
            self.assert_validation_failed(classified, output, "alignment")

    def test_validator_rejects_signature_date_alignment_and_indent(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            paragraph = paragraph_with_type(document, classified, "signature", 1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_first_line_indent_chars(paragraph, 2)
            document.save(output)
            result = validate_document(classified, output)
            checks = {error["check"] for error in result["errors"]}
            self.assertEqual(result["status"], "VALIDATION_FAILED")
            self.assertIn("alignment", checks)
            self.assertIn("first_line_indent", checks)

    def test_validator_rejects_body_space_after(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            paragraph_with_type(
                document, classified, "body"
            ).paragraph_format.space_after = Pt(10)
            document.save(output)
            self.assert_validation_failed(classified, output, "paragraph_spacing")

    def test_validator_rejects_heading_space_before(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            paragraph_with_type(
                document, classified, "heading_1"
            ).paragraph_format.space_before = Pt(6)
            document.save(output)
            self.assert_validation_failed(classified, output, "paragraph_spacing")

    def test_validator_rejects_normal_style_spacing(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            document.styles["Normal"].paragraph_format.space_after = Pt(10)
            document.save(output)
            self.assert_validation_failed(classified, output, "normal_style_spacing")

    def test_validator_rejects_document_default_spacing(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            spacing = document.styles.element.find(
                f"{qn('w:docDefaults')}/{qn('w:pPrDefault')}/"
                f"{qn('w:pPr')}/{qn('w:spacing')}"
            )
            spacing.set(qn("w:after"), "200")
            document.save(output)
            self.assert_validation_failed(classified, output, "default_style_spacing")

    def test_validator_rejects_wrong_signature_font(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            paragraph = paragraph_with_text(document, "XX部门")
            set_all_run_fonts(paragraph.runs[0], "宋体")
            document.save(output)
            self.assert_validation_failed(classified, output, "font_eastAsia")

    def test_validator_rejects_wrong_digit_font(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            classified, output, _, _ = build_output(Path(directory))
            document = Document(output)
            paragraph = paragraph_with_text(document, "关于2026年度工作安排的通知")
            digit_run = next(run for run in paragraph.runs if run.text == "2026")
            set_all_run_fonts(digit_run, "仿宋_GB2312")
            document.save(output)
            self.assert_validation_failed(classified, output, "font_eastAsia")


if __name__ == "__main__":
    unittest.main()
