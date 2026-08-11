from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SAMPLE_TEXTS = [
    "关于进一步加强人工智能应用管理工作的通知",
    "一、总体要求",
    "这是一级标题后的正文内容，不允许发生任何修改。",
    "",
    "（一）主要任务",
    "这是二级标题后的正文内容，不允许发生任何修改。",
    "1. 加强组织管理",
    "这是三级标题后的正文内容，不允许发生任何修改。",
    "（1）明确责任分工",
    "这是四级标题后的正文内容，不允许发生任何修改。",
    "附件：1.人工智能应用管理任务表",
    "　　　2.相关工作说明",
]


def create_messy_docx(path: str | Path) -> Path:
    destination = Path(path)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(9)
    section.page_height = Inches(12)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(1.0)

    for text in SAMPLE_TEXTS:
        paragraph = document.add_paragraph()
        if text:
            run = paragraph.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            fonts.set(qn("w:eastAsia"), "宋体")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.line_spacing = Pt(12)
    document.save(destination)
    return destination


def nonempty_texts(path: str | Path) -> list[str]:
    return [paragraph.text for paragraph in Document(path).paragraphs if paragraph.text != ""]
