from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from scripts.classify import classify_document
from scripts.extract import extract_document
from scripts.render_docx import render_document
from scripts.validate import validate_document
from tests.support import create_messy_docx, nonempty_texts


ROOT = Path(__file__).resolve().parents[1]
TEMP_ROOT = ROOT / ".tmp" / "tests"


class RenderTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        TEMP_ROOT.mkdir(parents=True, exist_ok=True)

    def test_render_and_real_docx_validation(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            folder = Path(directory)
            source = create_messy_docx(folder / "混乱输入.docx")
            extracted = extract_document(source)
            classified, analysis = classify_document(extracted)
            self.assertEqual(analysis["needs_review"], [])

            output = folder / "规范输出.docx"
            render_document(classified, output)
            result = validate_document(classified, output)
            self.assertTrue(result["success"], result["errors"])
            self.assertEqual(nonempty_texts(source), nonempty_texts(output))

            rendered = Document(output)
            self.assertEqual(len(rendered.paragraphs), 14)
            self.assertEqual(rendered.paragraphs[1].text, "")
            self.assertEqual(rendered.paragraphs[-3].text, "")
            source_section = Document(source).sections[0]
            output_section = rendered.sections[0]
            self.assertEqual(output_section.page_width, source_section.page_width)
            self.assertEqual(output_section.page_height, source_section.page_height)
            self.assertEqual(output_section.left_margin, source_section.left_margin)
            self.assertEqual(output_section.right_margin, source_section.right_margin)

            formats = result["validation"]["observed_formats"]
            self.assertEqual(formats["title"]["eastAsia"], ["方正小标宋简体"])
            self.assertEqual(formats["title"]["size_pt"], [22.0])
            self.assertEqual(formats["heading_1"]["eastAsia"], ["黑体"])
            self.assertEqual(formats["heading_2"]["eastAsia"], ["楷体_GB2312"])
            self.assertEqual(formats["heading_3"]["eastAsia"], ["仿宋_GB2312"])
            self.assertEqual(formats["heading_4"]["eastAsia"], ["仿宋_GB2312"])
            self.assertEqual(formats["body"]["size_pt"], [16.0])
            self.assertEqual(formats["attachment"]["line_spacing_pt"], [30.0])

    def test_renderer_refuses_to_overwrite_source(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            source = create_messy_docx(Path(directory) / "source.docx")
            classified, _ = classify_document(extract_document(source))
            with self.assertRaisesRegex(Exception, "must not overwrite"):
                render_document(classified, source)

    def test_validator_fails_closed_after_text_tampering(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            folder = Path(directory)
            source = create_messy_docx(folder / "source.docx")
            classified, _ = classify_document(extract_document(source))
            output = folder / "output.docx"
            render_document(classified, output)
            tampered = Document(output)
            tampered.paragraphs[0].runs[0].text += "篡改"
            tampered.save(output)
            result = validate_document(classified, output)
            self.assertFalse(result["success"])
            self.assertEqual(result["code"], "VALIDATION_FAILED")
            self.assertIsNone(result["output_file"])
            self.assertIn("text_integrity", {error["check"] for error in result["errors"]})


if __name__ == "__main__":
    unittest.main()
