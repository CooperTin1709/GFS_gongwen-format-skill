from __future__ import annotations

import base64
import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION

from scripts.extract import preflight_docx
from scripts.utils import PipelineError
from tests.support import create_messy_docx, nonempty_texts


ROOT = Path(__file__).resolve().parents[1]
TEMP_ROOT = ROOT / ".tmp" / "tests"
MAIN = ROOT / "scripts" / "main.py"


def run_cli(*arguments: object) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, '-X', 'utf8', str(MAIN), *(str(argument) for argument in arguments)],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
        capture_output=True,
        check=False,
    )


class EndToEndTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        TEMP_ROOT.mkdir(parents=True, exist_ok=True)

    def test_real_cli_pipeline_and_idempotence(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            folder = Path(directory)
            source = create_messy_docx(folder / "输入.docx")
            work_1 = folder / "work-1"
            output_1 = folder / "输出一.docx"

            analyze = run_cli("analyze", source, "--work-dir", work_1)
            self.assertEqual(analyze.returncode, 0, analyze.stderr + analyze.stdout)
            analysis = json.loads((work_1 / "analysis.json").read_text(encoding="utf-8"))
            self.assertEqual(analysis["needs_review"], [])

            render = run_cli("render", work_1 / "document.json", "--output", output_1)
            self.assertEqual(render.returncode, 0, render.stderr + render.stdout)
            first_result = json.loads((folder / "result.json").read_text(encoding="utf-8"))
            self.assertTrue(first_result["success"], first_result["errors"])

            independent = run_cli(
                "validate",
                work_1 / "document.json",
                "--output",
                output_1,
                "--result-file",
                folder / "independent-validation.json",
            )
            self.assertEqual(independent.returncode, 0, independent.stderr + independent.stdout)
            independent_result = json.loads(
                (folder / "independent-validation.json").read_text(encoding="utf-8")
            )
            self.assertTrue(independent_result["validation"]["text_identical"])
            self.assertTrue(independent_result["validation"]["blank_after_title"])
            self.assertTrue(independent_result["validation"]["blank_before_attachment"])

            work_2 = folder / "work-2"
            output_2 = folder / "输出二.docx"
            second = run_cli(
                "format",
                output_1,
                "--output",
                output_2,
                "--work-dir",
                work_2,
            )
            self.assertEqual(second.returncode, 0, second.stderr + second.stdout)
            second_result = json.loads((work_2 / "result.json").read_text(encoding="utf-8"))
            self.assertTrue(second_result["success"], second_result["errors"])
            self.assertEqual(nonempty_texts(source), nonempty_texts(output_1))
            self.assertEqual(nonempty_texts(output_1), nonempty_texts(output_2))
            self.assertEqual(
                [paragraph.text for paragraph in Document(output_1).paragraphs],
                [paragraph.text for paragraph in Document(output_2).paragraphs],
            )
            self.assertEqual(
                first_result["validation"]["observed_formats"],
                second_result["validation"]["observed_formats"],
            )

    def test_needs_review_blocks_format(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            folder = Path(directory)
            document = Document()
            document.add_paragraph("工作安排")
            source = folder / "ambiguous.docx"
            document.save(source)
            result = run_cli("format", source, "--output", folder / "out.docx", "--work-dir", folder / "work")
            self.assertEqual(result.returncode, 3)
            report = json.loads((folder / "work" / "result.json").read_text(encoding="utf-8"))
            self.assertEqual(report["code"], "NEEDS_REVIEW")
            self.assertFalse((folder / "out.docx").exists())


class PreflightTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        TEMP_ROOT.mkdir(parents=True, exist_ok=True)

    def assert_code(self, path: Path, expected: str) -> None:
        with self.assertRaises(PipelineError) as context:
            preflight_docx(path)
        self.assertEqual(context.exception.code, expected)

    def test_non_docx(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            path = Path(directory) / "input.txt"
            path.write_text("not a docx", encoding="utf-8")
            self.assert_code(path, "INVALID_INPUT")

    def test_corrupt_docx(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            path = Path(directory) / "broken.docx"
            path.write_bytes(b"not a zip package")
            self.assert_code(path, "INVALID_DOCX")

    def test_table_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            path = Path(directory) / "table.docx"
            document = Document()
            document.add_table(rows=1, cols=1).cell(0, 0).text = "敏感表格内容"
            document.save(path)
            self.assert_code(path, "UNSUPPORTED_COMPLEX_CONTENT")

    def test_image_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            folder = Path(directory)
            png = folder / "pixel.png"
            png.write_bytes(
                base64.b64decode(
                    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9Zl1sAAAAASUVORK5CYII="
                )
            )
            path = folder / "image.docx"
            document = Document()
            document.add_picture(str(png))
            document.save(path)
            self.assert_code(path, "UNSUPPORTED_COMPLEX_CONTENT")

    def test_multiple_sections_are_rejected(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            path = Path(directory) / "sections.docx"
            document = Document()
            document.add_paragraph("第一节")
            document.add_section(WD_SECTION.NEW_PAGE)
            document.add_paragraph("第二节")
            document.save(path)
            self.assert_code(path, "UNSUPPORTED_COMPLEX_CONTENT")


if __name__ == "__main__":
    unittest.main()
