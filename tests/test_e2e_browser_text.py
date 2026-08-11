from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

from scripts.input_adapter import adapt_source_text


ROOT = Path(__file__).resolve().parents[1]
TEMP_ROOT = ROOT / ".tmp" / "tests"
MAIN = ROOT / "scripts" / "main.py"
SAMPLE = ROOT / "samples" / "browser_input.txt"
AMBIGUOUS_TEXT = """关于测试工作的通知
一、总体要求
1. 这是长度处于灰区且没有明确句末标记的模糊编号段落内容
后续正文。"""


def run_cli(*arguments: object) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, "-X", "utf8", str(MAIN), *(str(item) for item in arguments)],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
        capture_output=True,
        check=False,
    )


class BrowserTextEndToEndTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        TEMP_ROOT.mkdir(parents=True, exist_ok=True)

    def test_browser_text_cli_success_and_idempotence(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            folder = Path(directory)
            output_1 = folder / "run-1"
            output_2 = folder / "run-2"

            first = run_cli("--text-file", SAMPLE, "--output-dir", output_1)
            self.assertEqual(first.returncode, 0, first.stderr + first.stdout)
            first_summary = json.loads(first.stdout)
            self.assertEqual(first_summary["status"], "SUCCESS")
            self.assertTrue(first_summary["validation_passed"])
            self.assertNotIn("人工智能技术快速发展", first.stdout)

            second = run_cli("--text-file", SAMPLE, "--output-dir", output_2)
            self.assertEqual(second.returncode, 0, second.stderr + second.stdout)
            self.assertEqual(json.loads(second.stdout)["status"], "SUCCESS")

            result_1 = json.loads((output_1 / "result.json").read_text(encoding="utf-8"))
            result_2 = json.loads((output_2 / "result.json").read_text(encoding="utf-8"))
            self.assertTrue(result_1["validation"]["text_identical"])
            self.assertEqual(result_1["verified_against"], "browser_extracted_text")
            self.assertNotIn("unknown", result_1["classification_counts"])

            source_text = SAMPLE.read_text(encoding="utf-8")
            expected = [item["text"] for item in adapt_source_text(source_text)["paragraphs"]]
            document_1 = Document(output_1 / "formatted.docx")
            document_2 = Document(output_2 / "formatted.docx")
            actual_1 = [paragraph.text for paragraph in document_1.paragraphs if paragraph.text != ""]
            actual_2 = [paragraph.text for paragraph in document_2.paragraphs if paragraph.text != ""]
            self.assertEqual(actual_1, expected)
            self.assertEqual(actual_2, expected)
            self.assertEqual(
                [paragraph.text for paragraph in document_1.paragraphs],
                [paragraph.text for paragraph in document_2.paragraphs],
            )
            self.assertEqual(
                result_1["validation"]["observed_formats"],
                result_2["validation"]["observed_formats"],
            )


class NeedsReviewSimulationTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        TEMP_ROOT.mkdir(parents=True, exist_ok=True)

    def make_text_file(self, folder: Path) -> Path:
        text_file = folder / "ambiguous.txt"
        text_file.write_text(AMBIGUOUS_TEXT, encoding="utf-8")
        return text_file

    def test_valid_override_resumes_once(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            folder = Path(directory)
            text_file = self.make_text_file(folder)
            output = folder / "output"
            first = run_cli("--text-file", text_file, "--output-dir", output)
            self.assertEqual(first.returncode, 3, first.stderr + first.stdout)
            summary = json.loads(first.stdout)
            self.assertEqual(summary["status"], "NEEDS_REVIEW")
            self.assertEqual(summary["review_count"], 1)
            review = json.loads((output / "review.json").read_text(encoding="utf-8"))
            self.assertEqual(review[0]["candidate_types"], ["heading_3", "body"])
            self.assertEqual(
                set(review[0]),
                {"id", "text", "previous_text", "next_text", "candidate_types", "suggested_type"},
            )

            overrides = folder / "overrides.json"
            overrides.write_text(
                json.dumps({review[0]["id"]: "heading_3"}, ensure_ascii=False),
                encoding="utf-8",
            )
            resumed = run_cli(
                "--text-file",
                text_file,
                "--output-dir",
                output,
                "--overrides",
                overrides,
            )
            self.assertEqual(resumed.returncode, 0, resumed.stderr + resumed.stdout)
            self.assertEqual(json.loads(resumed.stdout)["status"], "SUCCESS")
            self.assertTrue((output / "formatted.docx").is_file())

    def test_invalid_candidate_type_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            folder = Path(directory)
            text_file = self.make_text_file(folder)
            overrides = folder / "overrides.json"
            overrides.write_text('{"p0003":"heading_1"}', encoding="utf-8")
            result = run_cli(
                "--text-file",
                text_file,
                "--output-dir",
                folder / "output",
                "--overrides",
                overrides,
            )
            self.assertEqual(result.returncode, 2)
            self.assertEqual(json.loads(result.stdout)["status"], "INVALID_OVERRIDE")

    def test_unknown_override_id_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            folder = Path(directory)
            text_file = self.make_text_file(folder)
            overrides = folder / "overrides.json"
            overrides.write_text('{"p9999":"heading_3"}', encoding="utf-8")
            result = run_cli(
                "--text-file",
                text_file,
                "--output-dir",
                folder / "output",
                "--overrides",
                overrides,
            )
            self.assertEqual(result.returncode, 2)
            self.assertEqual(json.loads(result.stdout)["status"], "INVALID_OVERRIDE")

    def test_malformed_override_json_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory(dir=TEMP_ROOT) as directory:
            folder = Path(directory)
            text_file = self.make_text_file(folder)
            overrides = folder / "overrides.json"
            overrides.write_text('{"p0003":', encoding="utf-8")
            result = run_cli(
                "--text-file",
                text_file,
                "--output-dir",
                folder / "output",
                "--overrides",
                overrides,
            )
            self.assertEqual(result.returncode, 2)
            self.assertEqual(json.loads(result.stdout)["status"], "INVALID_OVERRIDE")


if __name__ == "__main__":
    unittest.main()
