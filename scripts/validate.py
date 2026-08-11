from __future__ import annotations

from collections import defaultdict
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from .docx_utils import (
        DIGIT_RUN_RE,
        FONT_ATTRIBUTES,
        font_values,
        indent_values,
        is_digit_run_text,
        line_spacing,
        split_text_by_digit_runs,
    )
    from .render_docx import build_render_plan, resolve_classifications
    from .utils import PipelineError, load_rules, resolved_format_rule, workspace_path
except ImportError:  # pragma: no cover - direct script import fallback
    from docx_utils import (
        DIGIT_RUN_RE,
        FONT_ATTRIBUTES,
        font_values,
        indent_values,
        is_digit_run_text,
        line_spacing,
        split_text_by_digit_runs,
    )
    from render_docx import build_render_plan, resolve_classifications
    from utils import PipelineError, load_rules, resolved_format_rule, workspace_path


ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def _first_mismatch(expected: list[str], actual: list[str]) -> int | None:
    for index, (left, right) in enumerate(zip(expected, actual)):
        if left != right:
            return index
    return min(len(expected), len(actual)) if len(expected) != len(actual) else None


def _blank_count_after(paragraphs: list[Any], index: int) -> int:
    count = 0
    for paragraph in paragraphs[index + 1 :]:
        if paragraph.text != "":
            break
        count += 1
    return count


def _blank_count_before(paragraphs: list[Any], index: int) -> int:
    count = 0
    for paragraph in reversed(paragraphs[:index]):
        if paragraph.text != "":
            break
        count += 1
    return count


def _indent_chars(values: dict[str, str | None]) -> float | None:
    raw = values["firstLineChars"]
    if raw is None:
        return 0.0
    try:
        return int(raw) / 100
    except ValueError:
        return None


def validate_document(
    document_data: dict[str, Any],
    output: str | Path,
    *,
    overrides: dict[str, str] | None = None,
    rules_path: str | Path | None = None,
    workspace: str | Path | None = None,
) -> dict[str, Any]:
    output_path = workspace_path(output, must_exist=True, workspace=workspace)
    rules = load_rules(rules_path)
    resolved = resolve_classifications(document_data, overrides)
    plan = build_render_plan(resolved, rules["blank_policy"])
    try:
        rendered = Document(output_path)
    except Exception as exc:
        raise PipelineError(
            "VALIDATION_FAILED",
            "Rendered DOCX cannot be reopened.",
            details={"reason": type(exc).__name__},
        ) from exc

    errors: list[dict[str, Any]] = []
    expected_nonempty = [item["text"] for item in document_data["paragraphs"]]
    actual_nonempty = [
        paragraph.text for paragraph in rendered.paragraphs if paragraph.text != ""
    ]
    text_mismatch = _first_mismatch(expected_nonempty, actual_nonempty)
    if text_mismatch is not None:
        errors.append(
            {
                "check": "text_integrity",
                "message": "Output non-empty text differs from Browser-extracted text.",
                "first_mismatch_index": text_mismatch,
                "expected_count": len(expected_nonempty),
                "actual_count": len(actual_nonempty),
            }
        )

    expected_all = [item["text"] for item in plan]
    actual_all = [paragraph.text for paragraph in rendered.paragraphs]
    sequence_mismatch = _first_mismatch(expected_all, actual_all)
    if sequence_mismatch is not None:
        errors.append(
            {
                "check": "paragraph_sequence",
                "message": "Rendered paragraphs do not match the canonical blank policy.",
                "first_mismatch_index": sequence_mismatch,
                "expected_count": len(expected_all),
                "actual_count": len(actual_all),
            }
        )

    observations: dict[str, dict[str, set[Any]]] = defaultdict(
        lambda: {
            "eastAsia": set(),
            "ascii": set(),
            "hAnsi": set(),
            "cs": set(),
            "size_pt": set(),
            "line_spacing_pt": set(),
            "alignment": set(),
            "first_line_indent_chars": set(),
        }
    )
    expected_line = float(rules["global"]["line_spacing_pt"])
    digit_font = str(rules["global"]["digit_font"])

    for index, item in enumerate(plan[: len(rendered.paragraphs)]):
        paragraph = rendered.paragraphs[index]
        paragraph_type = item["classification"]
        line_rule, line_points = line_spacing(paragraph)
        observations[paragraph_type]["line_spacing_pt"].add(line_points)
        observations[paragraph_type]["alignment"].add(
            getattr(paragraph.alignment, "name", None)
        )
        if line_rule != "exact" or line_points != expected_line:
            errors.append(
                {
                    "check": "line_spacing",
                    "paragraph_id": item["id"],
                    "expected_rule": "exact",
                    "expected_pt": expected_line,
                    "actual_rule": line_rule,
                    "actual_pt": line_points,
                }
            )

        rule = (
            resolved_format_rule(rules, paragraph_type)
            if paragraph_type != "blank"
            else None
        )
        expected_indent = (
            float(rule["first_line_indent_chars"]) if rule is not None else 0.0
        )
        actual_indent_values = indent_values(paragraph)
        actual_indent = _indent_chars(actual_indent_values)
        observations[paragraph_type]["first_line_indent_chars"].add(actual_indent)
        conflicts = {
            key: actual_indent_values[key]
            for key in ("firstLine", "hanging", "hangingChars")
            if actual_indent_values[key] is not None
        }
        if actual_indent != expected_indent or conflicts:
            errors.append(
                {
                    "check": "first_line_indent",
                    "paragraph_id": item["id"],
                    "expected_chars": expected_indent,
                    "actual_chars": actual_indent,
                    "actual_firstLineChars": actual_indent_values["firstLineChars"],
                    "conflicts": conflicts,
                }
            )

        if paragraph_type == "blank":
            continue
        assert rule is not None

        expected_segments = split_text_by_digit_runs(str(item["text"]))
        text_runs = [run for run in paragraph.runs if run.text != ""]
        actual_segments = [
            (run.text, is_digit_run_text(run.text)) for run in text_runs
        ]
        if actual_segments != expected_segments:
            errors.append(
                {
                    "check": "run_structure",
                    "paragraph_id": item["id"],
                    "message": "Runs must preserve text and isolate consecutive digits.",
                }
            )

        for run in text_runs:
            digit_only = is_digit_run_text(run.text)
            contains_digit = DIGIT_RUN_RE.search(run.text) is not None
            expected_font = digit_font if digit_only else str(rule["font"])
            actual_fonts = font_values(run)
            for attribute, actual_font in actual_fonts.items():
                observations[paragraph_type][attribute].add(actual_font)
                if actual_font != expected_font:
                    errors.append(
                        {
                            "check": f"font_{attribute}",
                            "paragraph_id": item["id"],
                            "run_text": run.text,
                            "expected": expected_font,
                            "actual": actual_font,
                        }
                    )
            if contains_digit and not digit_only:
                errors.append(
                    {
                        "check": "digit_font_run_content",
                        "paragraph_id": item["id"],
                        "message": "A run containing digits also contains non-digit characters.",
                    }
                )
            if any(value == digit_font for value in actual_fonts.values()) and not digit_only:
                errors.append(
                    {
                        "check": "digit_font_run_content",
                        "paragraph_id": item["id"],
                        "message": "digit font run contains non-digit characters",
                    }
                )

            size_pt = run.font.size.pt if run.font.size is not None else None
            observations[paragraph_type]["size_pt"].add(size_pt)
            if size_pt != float(rule["size_pt"]):
                errors.append(
                    {
                        "check": "font_size",
                        "paragraph_id": item["id"],
                        "expected_pt": float(rule["size_pt"]),
                        "actual_pt": size_pt,
                    }
                )
            if "bold" in rule and run.font.bold != bool(rule["bold"]):
                errors.append(
                    {
                        "check": "bold",
                        "paragraph_id": item["id"],
                        "expected": bool(rule["bold"]),
                        "actual": run.font.bold,
                    }
                )

        expected_alignment = rule.get("alignment")
        if (
            expected_alignment in ALIGNMENTS
            and paragraph.alignment != ALIGNMENTS[expected_alignment]
        ):
            errors.append(
                {
                    "check": "alignment",
                    "paragraph_id": item["id"],
                    "expected": expected_alignment,
                    "actual": getattr(paragraph.alignment, "name", None),
                }
            )

    title_index = next(
        (index for index, item in enumerate(plan) if item["classification"] == "title"),
        None,
    )
    title_blank_count = (
        _blank_count_after(rendered.paragraphs, title_index)
        if title_index is not None
        else 0
    )
    title_blank_ok = title_index is None or title_blank_count == int(
        rules["blank_policy"]["after_title"]
    )
    if not title_blank_ok:
        errors.append(
            {
                "check": "blank_after_title",
                "actual": title_blank_count,
                "expected": int(rules["blank_policy"]["after_title"]),
            }
        )

    attachment_index = next(
        (index for index, item in enumerate(plan) if item["classification"] == "attachment"),
        None,
    )
    attachment_blank_count = (
        _blank_count_before(rendered.paragraphs, attachment_index)
        if attachment_index is not None
        else 0
    )
    attachment_blank_ok = attachment_index is None or attachment_blank_count == int(
        rules["blank_policy"]["before_attachment"]
    )
    if not attachment_blank_ok:
        errors.append(
            {
                "check": "blank_before_attachment",
                "actual": attachment_blank_count,
                "expected": int(rules["blank_policy"]["before_attachment"]),
            }
        )

    observed_formats = {
        key: {
            field: sorted(values, key=lambda value: (value is None, str(value)))
            for field, values in fields.items()
        }
        for key, fields in sorted(observations.items())
    }
    success = not errors
    return {
        "status": "SUCCESS" if success else "VALIDATION_FAILED",
        "source_type": "browser_text",
        "output_file": str(output_path) if success else None,
        "validation_passed": success,
        "verified_against": "browser_extracted_text",
        "errors": errors,
        "validation": {
            "text_identical": text_mismatch is None,
            "paragraph_sequence_matches_plan": sequence_mismatch is None,
            "source_nonempty_count": len(expected_nonempty),
            "output_nonempty_count": len(actual_nonempty),
            "blank_after_title": title_blank_ok,
            "blank_after_title_count": title_blank_count,
            "blank_before_attachment": attachment_blank_ok,
            "blank_before_attachment_count": attachment_blank_count,
            "observed_formats": observed_formats,
        },
    }
