from __future__ import annotations

import datetime as dt
import zipfile
from pathlib import Path
from typing import Any

from docx import Document

try:
    from .utils import PipelineError, workspace_path, write_json
except ImportError:  # pragma: no cover - direct script import fallback
    from utils import PipelineError, workspace_path, write_json


def _has_xpath(element: Any, expression: str) -> bool:
    return bool(element.xpath(expression))


def _header_footer_has_content(container: Any) -> bool:
    if any(paragraph.text != "" for paragraph in container.paragraphs):
        return True
    if container.tables:
        return True
    element = container._element
    return any(
        _has_xpath(element, expression)
        for expression in (".//w:drawing", ".//w:pict", ".//w:txbxContent")
    )


def preflight_docx(
    source: str | Path,
    *,
    workspace: str | Path | None = None,
) -> tuple[Any, dict[str, Any]]:
    path = workspace_path(source, must_exist=True, workspace=workspace)
    if not path.is_file() or path.suffix.lower() != ".docx":
        raise PipelineError(
            "INVALID_INPUT",
            "Input must be an existing .docx file.",
            details={"path": str(path)},
        )
    if not zipfile.is_zipfile(path):
        raise PipelineError(
            "INVALID_DOCX",
            "The input is not a valid DOCX package.",
            details={"path": str(path)},
        )
    try:
        document = Document(path)
    except Exception as exc:
        raise PipelineError(
            "INVALID_DOCX",
            "The DOCX package cannot be opened.",
            details={"path": str(path), "reason": type(exc).__name__},
        ) from exc

    body = document._element.body
    findings: list[str] = []
    if len(document.sections) != 1:
        findings.append("multiple_sections")
    if document.tables or _has_xpath(body, ".//w:tbl"):
        findings.append("tables")
    if _has_xpath(body, ".//w:drawing") or _has_xpath(body, ".//w:pict"):
        findings.append("drawings_or_images")
    if _has_xpath(body, ".//w:txbxContent"):
        findings.append("text_boxes")
    if _has_xpath(body, ".//w:object") or _has_xpath(body, ".//w:altChunk"):
        findings.append("embedded_or_alt_content")
    with zipfile.ZipFile(path) as package:
        if any(name.startswith("word/media/") for name in package.namelist()):
            findings.append("media")

    for section in document.sections:
        containers = (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        )
        if any(_header_footer_has_content(container) for container in containers):
            findings.append("header_or_footer_content")
            break

    findings = sorted(set(findings))
    if findings:
        raise PipelineError(
            "UNSUPPORTED_COMPLEX_CONTENT",
            "This version only rebuilds plain-text, single-section DOCX files.",
            details={"unsupported": findings},
        )

    return document, {
        "supported": True,
        "section_count": len(document.sections),
        "table_count": len(document.tables),
        "inline_shape_count": len(document.inline_shapes),
        "complex_content": [],
    }


def _length_value(value: Any) -> int | None:
    return None if value is None else int(value)


def _section_metadata(section: Any) -> dict[str, Any]:
    orientation = section.orientation
    orientation_name = getattr(orientation, "name", str(orientation)).lower()
    return {
        "page_width": _length_value(section.page_width),
        "page_height": _length_value(section.page_height),
        "top_margin": _length_value(section.top_margin),
        "bottom_margin": _length_value(section.bottom_margin),
        "left_margin": _length_value(section.left_margin),
        "right_margin": _length_value(section.right_margin),
        "header_distance": _length_value(section.header_distance),
        "footer_distance": _length_value(section.footer_distance),
        "gutter": _length_value(section.gutter),
        "orientation": orientation_name,
    }


def extract_document(
    source: str | Path,
    *,
    workspace: str | Path | None = None,
) -> dict[str, Any]:
    path = workspace_path(source, must_exist=True, workspace=workspace)
    document, preflight = preflight_docx(path, workspace=workspace)
    paragraphs: list[dict[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text
        alignment = paragraph.alignment
        paragraphs.append(
            {
                "id": f"p{index + 1:04d}",
                "index": index,
                "text": text,
                "is_blank": text == "",
                "original_style": {
                    "style_name": paragraph.style.name if paragraph.style else None,
                    "alignment": getattr(alignment, "name", None),
                    "run_count": len(paragraph.runs),
                },
                "classification": None,
                "confidence": None,
                "classification_source": None,
            }
        )
    return {
        "schema_version": 1,
        "source_file": str(path),
        "paragraphs": paragraphs,
        "preflight": preflight,
        "metadata": {
            "source_type": "docx",
            "extracted_at": dt.datetime.now(dt.timezone.utc).isoformat(),
            "paragraph_count": len(paragraphs),
            "section": _section_metadata(document.sections[0]),
        },
    }


def markdown_view(document_data: dict[str, Any]) -> str:
    lines = [
        "# DOCX paragraph view",
        "",
        "> This view is for classification only. The JSON `text` fields remain authoritative.",
        "",
    ]
    for paragraph in document_data["paragraphs"]:
        lines.append(f"<!-- {paragraph['id']} -->")
        if paragraph["is_blank"]:
            lines.append("[BLANK]")
        else:
            for line in paragraph["text"].split("\n"):
                lines.append(f"    {line}")
        lines.append("")
    return "\n".join(lines)


def write_extraction(document_data: dict[str, Any], work_dir: str | Path) -> dict[str, str]:
    destination = workspace_path(work_dir)
    destination.mkdir(parents=True, exist_ok=True)
    json_path = write_json(destination / "document.json", document_data)
    markdown_path = destination / "document.md"
    markdown_path.write_text(markdown_view(document_data), encoding="utf-8")
    return {"document_json": str(json_path), "document_md": str(markdown_path)}
