from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from .docx_utils import (
        set_exact_line_spacing,
        set_first_line_indent_chars,
        set_run_format,
        split_text_by_digit_runs,
    )
    from .utils import (
        PARAGRAPH_TYPES,
        PipelineError,
        load_json,
        load_rules,
        resolved_format_rule,
        workspace_path,
    )
except ImportError:  # pragma: no cover - direct script import fallback
    from docx_utils import (
        set_exact_line_spacing,
        set_first_line_indent_chars,
        set_run_format,
        split_text_by_digit_runs,
    )
    from utils import (
        PARAGRAPH_TYPES,
        PipelineError,
        load_json,
        load_rules,
        resolved_format_rule,
        workspace_path,
    )


def load_overrides(path: str | Path | None) -> dict[str, str]:
    if path is None:
        return {}
    try:
        override_path = workspace_path(path, must_exist=True)
        data = load_json(override_path)
    except PipelineError as exc:
        raise PipelineError(
            "INVALID_OVERRIDE",
            "Unable to read a valid overrides JSON file.",
            details=exc.details,
        ) from exc
    if not isinstance(data, dict) or not all(
        isinstance(key, str) and isinstance(value, str) for key, value in data.items()
    ):
        raise PipelineError(
            "INVALID_OVERRIDE",
            "Overrides must be a JSON object mapping paragraph id to type.",
        )
    return data


def resolve_classifications(
    document_data: dict[str, Any], overrides: dict[str, str] | None = None
) -> list[dict[str, Any]]:
    paragraphs = document_data.get("paragraphs")
    if not isinstance(paragraphs, list):
        raise PipelineError("INVALID_INPUT", "Canonical paragraph data is missing.")

    override_map = overrides or {}
    known_ids = {item.get("id") for item in paragraphs}
    unknown_override_ids = sorted(set(override_map) - known_ids)
    if unknown_override_ids:
        raise PipelineError(
            "INVALID_OVERRIDE",
            "Override contains unknown paragraph ids.",
            details={"ids": unknown_override_ids},
        )

    resolved: list[dict[str, Any]] = []
    unresolved: list[str] = []
    for item in paragraphs:
        paragraph = dict(item)
        paragraph_id = str(paragraph.get("id"))
        paragraph_type = paragraph.get("classification")
        if paragraph_id in override_map:
            if paragraph_type != "unknown":
                raise PipelineError(
                    "INVALID_OVERRIDE",
                    "Only NEEDS_REVIEW paragraphs may be overridden.",
                    details={"id": paragraph_id},
                )
            candidate_types = paragraph.get("candidate_types", [])
            override_type = override_map[paragraph_id]
            if override_type not in candidate_types:
                raise PipelineError(
                    "INVALID_OVERRIDE",
                    "Override type is not one of the paragraph candidate_types.",
                    details={
                        "id": paragraph_id,
                        "type": override_type,
                        "candidate_types": candidate_types,
                    },
                )
            paragraph["classification"] = override_type
            paragraph["classification_source"] = "override"
            paragraph["confidence"] = 1.0
            paragraph_type = override_type

        if paragraph_type == "unknown":
            unresolved.append(paragraph_id)
            continue
        if paragraph_type not in PARAGRAPH_TYPES:
            raise PipelineError(
                "INVALID_INPUT",
                "Canonical data contains an invalid classification.",
                details={"id": paragraph_id, "type": paragraph_type},
            )
        resolved.append(paragraph)

    if unresolved:
        if overrides is None:
            raise PipelineError(
                "NEEDS_REVIEW",
                "Ambiguous paragraphs require one constrained review.",
                details={"paragraph_ids": unresolved},
            )
        raise PipelineError(
            "INVALID_OVERRIDE",
            "The single override submission did not resolve every review item.",
            details={"missing_ids": unresolved},
        )
    if sum(item["classification"] == "title" for item in resolved) > 1:
        raise PipelineError(
            "INVALID_OVERRIDE",
            "At most one paragraph may be classified as title.",
        )
    return resolved


def _blank_record(identifier: str) -> dict[str, Any]:
    return {
        "id": identifier,
        "index": None,
        "text": "",
        "classification": "blank",
        "classification_source": "blank_policy",
        "confidence": 1.0,
    }


def build_render_plan(
    paragraphs: list[dict[str, Any]], blank_policy: dict[str, Any]
) -> list[dict[str, Any]]:
    plan: list[dict[str, Any]] = []
    attachment_started = False
    for item in paragraphs:
        paragraph = dict(item)
        paragraph_type = paragraph["classification"]
        if paragraph_type == "attachment" and not attachment_started:
            while plan and plan[-1]["classification"] == "blank":
                plan.pop()
            plan.extend(
                _blank_record(f"__blank_before_attachment_{number + 1}")
                for number in range(int(blank_policy["before_attachment"]))
            )
            attachment_started = True
        plan.append(paragraph)
        if paragraph_type == "title":
            plan.extend(
                _blank_record(f"__blank_after_title_{number + 1}")
                for number in range(int(blank_policy["after_title"]))
            )
    return plan


def _remove_default_paragraph(document: Any) -> None:
    if document.paragraphs:
        paragraph = document.paragraphs[0]
        paragraph._element.getparent().remove(paragraph._element)


ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def render_document(
    document_data: dict[str, Any],
    output: str | Path,
    *,
    overrides: dict[str, str] | None = None,
    rules_path: str | Path | None = None,
    workspace: str | Path | None = None,
) -> tuple[Path, list[dict[str, Any]]]:
    output_path = workspace_path(output, workspace=workspace)
    if output_path.suffix.lower() != ".docx":
        raise PipelineError("RENDER_FAILED", "Output path must end with .docx.")

    rules = load_rules(rules_path)
    resolved = resolve_classifications(document_data, overrides)
    plan = build_render_plan(resolved, rules["blank_policy"])
    document = Document()
    _remove_default_paragraph(document)
    line_spacing_pt = float(rules["global"]["line_spacing_pt"])
    digit_font = str(rules["global"]["digit_font"])

    for item in plan:
        paragraph = document.add_paragraph()
        paragraph_type = item["classification"]
        if paragraph_type != "blank":
            rule = resolved_format_rule(rules, paragraph_type)
            text = str(item["text"])
            for segment, is_digit in split_text_by_digit_runs(text):
                run = paragraph.add_run(segment)
                set_run_format(
                    run,
                    font_name=digit_font if is_digit else str(rule["font"]),
                    size_pt=float(rule["size_pt"]),
                    bold=bool(rule["bold"]) if "bold" in rule else None,
                )
            if paragraph.text != text:
                raise PipelineError(
                    "RENDER_FAILED",
                    "Run splitting changed source paragraph text.",
                    details={"id": item["id"]},
                )
            alignment = rule.get("alignment")
            if alignment in ALIGNMENTS:
                paragraph.alignment = ALIGNMENTS[alignment]
            set_first_line_indent_chars(
                paragraph, float(rule["first_line_indent_chars"])
            )
        set_exact_line_spacing(paragraph, line_spacing_pt)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        document.save(output_path)
    except Exception as exc:
        raise PipelineError(
            "RENDER_FAILED",
            "Unable to save the rendered DOCX.",
            details={"path": str(output_path), "reason": type(exc).__name__},
        ) from exc
    return output_path, plan
